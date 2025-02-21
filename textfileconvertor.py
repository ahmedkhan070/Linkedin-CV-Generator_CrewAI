import os
from docx import Document

def convert_md_to_docx(md_file, output_docx="Formatted_CV.docx"):
    if not os.path.exists(md_file):
        print(f"Error: {md_file} not found!")
        return

    with open(md_file, "r", encoding="utf-8") as file:
        lines = file.readlines()

    # Remove markdown markers (```) if present
    if lines[0].strip().startswith("```"):
        lines = lines[1:]  # Remove first line
    if lines[-1].strip().startswith("```"):
        lines = lines[:-1]  # Remove last line

    doc = Document()

    for line in lines:
        line = line.strip()
        if not line:
            continue

        if line.startswith("# "):  # H1 -> Heading 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):  # H2 -> Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith("### "):  # H3 -> Heading 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith("**") and line.endswith("**"):  # Bold standalone lines
            para = doc.add_paragraph()
            para.add_run(line.replace("**", "")).bold = True
        elif "**" in line:  # Bold inline text
            parts = line.split("**")
            para = doc.add_paragraph()
            for i, part in enumerate(parts):
                run = para.add_run(part)
                if i % 2 == 1:
                    run.bold = True
        elif line.startswith("* "):  # Bullet points
            doc.add_paragraph(line[2:], style="List Bullet")
        else:
            doc.add_paragraph(line)  # Normal text

    doc.save(output_docx)
    return output_docx